import io
import time
import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types
from google.genai.errors import APIError
from pydantic import BaseModel, Field
from typing import List
from pdf2image import convert_from_bytes

# =====================================================================
# 1. APPLICATION INITIALIZATION & CONFIGURATION
# =====================================================================
st.set_page_config(
    page_title="Document Layout AI Translator",
    page_icon="🌐",
    layout="centered"
)

st.title("PDF-to-Word Visual AI Translator")
st.write("Translates complex documents while maintaining layout blocks, headers, and page structures.")

# Initialize Persistent Session State for Downloads (Prevents Nested Button Trap)
if "translated_docx_bytes" not in st.session_state:
    st.session_state.translated_docx_bytes = None
if "current_file_name" not in st.session_state:
    st.session_state.current_file_name = ""

# =====================================================================
# 2. MASTER SECURITY ACCESS GATE
# =====================================================================
MASTER_PASSWORD = st.secrets.get("APP_PASSWORD")

if not MASTER_PASSWORD:
    st.error("No APP_PASSWORD configured in st.secrets. Set one before deploying.")
    st.stop()

user_password = st.text_input("Enter Cloud Access Password:", type="password")

if user_password != MASTER_PASSWORD:
    if user_password:
        st.error("Invalid Credentials. Please check your password configuration.")
    st.info("🔒 Authorization required to boot the programmatic translation pipeline.")
    st.stop()

# =====================================================================
# 3. USER INTERFACE CONTROLS
# =====================================================================
uploaded_file = st.file_uploader("Upload Source Document (PDF Format)", type="pdf")
target_lang = st.selectbox(
    "Translate to:",
    ["English", "Spanish", "French", "German", "Japanese", "Italian", "Portuguese", "Chinese"]
)

# Reset session state if a brand new file is uploaded
if uploaded_file and uploaded_file.name != st.session_state.current_file_name:
    st.session_state.translated_docx_bytes = None
    st.session_state.current_file_name = uploaded_file.name

# =====================================================================
# 4. STRUCTURED DATA SCHEMAS FOR GEMINI
# =====================================================================
class ContentBlock(BaseModel):
    block_type: str = Field(
        ...,
        description="The layout component type: 'heading_1', 'heading_2', 'paragraph', or 'page_number'."
    )
    text: str = Field(
        ...,
        description="The strictly translated text belonging to this layout block component."
    )

class PageLayoutTranslation(BaseModel):
    blocks: List[ContentBlock] = Field(
        ...,
        description="Sequential list of all structural layout blocks translated from the page."
    )

# =====================================================================
# 5. TRANSLATION PIPELINE CORE ENGINE (UPDATED FOR VISUAL MULTIMODAL)
# =====================================================================
if uploaded_file and st.button("Execute Translation Blueprint"):
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])

    # Read file bytes for pdf2image
    file_bytes = uploaded_file.read()
    
    # Convert PDF pages to images so Gemini can visually "read" the Arabic text
    try:
        pages_images = convert_from_bytes(file_bytes)
        total_pages = len(pages_images)
    except Exception as e:
        st.error(f"Could not convert PDF to images. Is poppler installed? Error: {e}")
        st.stop()

    if total_pages == 0:
        st.error("This PDF appears to have no pages.")
        st.stop()

    doc = Document()
    progress_bar = st.progress(0)
    status_msg = st.empty()

    previous_page_context = "This is the first page. No prior context exists."
    pipeline_failed = False

    for idx, page_image in enumerate(pages_images):
        status_msg.text(f"Visually processing & translating page {idx + 1} of {total_pages}...")

        # Convert PIL Image to bytes for Gemini API
        img_byte_arr = io.BytesIO()
        page_image.save(img_byte_arr, format='JPEG')
        img_bytes = img_byte_arr.getvalue()

        # Update the prompt to tell Gemini to look at the image and ignore watermarks
        prompt = f"""
        You are an elite expert document layout translator. Your task is to look at the provided page image, 
        read the text (which may be complex Arabic script or poetry), and translate it into {target_lang}.

        CONTEXT FROM PREVIOUS PAGE:
        \"\"\"{previous_page_context}\"\"\"

        CRITICAL DIRECTIVES:
        1. IGNORE any background watermarks, website URLs, or repetitive header overlays that are not part of the main text/article body.
        2. Pay close attention to structural formatting (like poetry hemistichs/couplets or standard paragraphs).

        BLOCK ARCHITECTURE LAYOUT DIRECTIVES:
        1. heading_1: Used for top-level primary document title names.
        2. heading_2: Used for section subheadings or chapter dividers.
        3. paragraph: Standard continuous reading paragraphs or bullet records.
        4. page_number: Footers, headers, or isolated indicator numbers matching page coordinates.
        """

        translation_data = None
        for attempt in range(2):
            try:
                # Pass BOTH the image bytes and the text prompt to Gemini
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[
                        types.Part.from_bytes(data=img_bytes, mime_type="image/jpeg"),
                        prompt
                    ],
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=PageLayoutTranslation,
                        temperature=0.2
                    )
                )
                translation_data = PageLayoutTranslation.model_validate_json(response.text)
                break
            except Exception as e:
                if attempt == 0:
                    status_msg.warning(f"Page {idx + 1}: Error occurred. Retrying once... ({e})")
                    time.sleep(2)
                    continue
                st.error(f"Page {idx + 1} failed after retry. Stopping pipeline. Error: {e}")
                pipeline_failed = True
                break

        if pipeline_failed:
            break

        # Build next page's context
        context_candidates = [
            b.text for b in translation_data.blocks
            if b.block_type != "page_number" and b.text.strip()
        ]
        if context_candidates:
            previous_page_context = (
                "The previous page ended with this translation text: "
                + " ".join(context_candidates[-2:])
            )

        # =====================================================================
        # 6. THE BLUE STYLING WORD ENGINE
        # =====================================================================
        for block in translation_data.blocks:
            b_type = block.block_type
            text = block.text.strip()

            if not text:
                continue

            if b_type == "heading_1":
                h = doc.add_heading(text, level=1)
                for run in h.runs:
                    run.font.name = 'Calibri'

            elif b_type == "heading_2":
                h = doc.add_heading(text, level=2)
                for run in h.runs:
                    run.font.name = 'Calibri'

            elif b_type == "page_number":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 51, 204)  # Professional Crisp Blue
                run.bold = True

            else:  # Paragraph / Default
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

        progress_bar.progress((idx + 1) / total_pages)
        time.sleep(0.5)

    if not pipeline_failed:
        status_msg.success("All pages parsed and compiled successfully!")

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        st.session_state.translated_docx_bytes = docx_buffer.getvalue()
    else:
        st.session_state.translated_docx_bytes = None

# =====================================================================
# 7. INDEPENDENT PERSISTENT DOWNLOAD ACCESS CONTROL
# =====================================================================
if st.session_state.translated_docx_bytes is not None:
    st.write("---")
    st.success("🎉 Document Translation Built Successfully! (right click and select Resume if you get a permission error)")

    st.download_button(
        label="📥 Download Translated Word Document (.docx)",
        data=st.session_state.translated_docx_bytes,
        file_name=f"translated_{st.session_state.current_file_name.replace('.pdf', '.docx')}",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )